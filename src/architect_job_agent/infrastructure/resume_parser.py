"""Resume parsing for .docx and .pdf files."""
from __future__ import annotations

from pathlib import Path

from ..core.exceptions import ResumeParsingError
from ..core.logging import get_logger

logger = get_logger(__name__)


class ResumeParser:
    SUPPORTED = {".docx", ".pdf"}

    def parse(self, path: str | Path) -> str:
        p = Path(path)
        if not p.exists():
            raise ResumeParsingError(f"Resume not found: {p}")
        suffix = p.suffix.lower()
        if suffix not in self.SUPPORTED:
            raise ResumeParsingError(f"Unsupported resume format: {suffix}")
        try:
            if suffix == ".docx":
                return self._parse_docx(p)
            return self._parse_pdf(p)
        except ResumeParsingError:
            raise
        except Exception as exc:
            raise ResumeParsingError(f"Failed to parse {p.name}: {exc}") from exc

    @staticmethod
    def _parse_docx(path: Path) -> str:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise ResumeParsingError("python-docx not installed") from exc
        try:
            doc = Document(str(path))
        except Exception as exc:
            raise ResumeParsingError(f"Corrupted DOCX: {exc}") from exc
        chunks: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        chunks.append(text)
        text = "\n".join(chunks).strip()
        if not text:
            raise ResumeParsingError("DOCX contained no extractable text")
        return text

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:
            raise ResumeParsingError("pypdf not installed") from exc
        try:
            reader = PdfReader(str(path))
        except Exception as exc:
            raise ResumeParsingError(f"Corrupted PDF: {exc}") from exc
        pages: list[str] = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:
                continue
        text = "\n".join(pages).strip()
        if not text:
            raise ResumeParsingError("PDF contained no extractable text")
        return text


class ResumeWriter:
    """Persists tailored resumes to disk as plain .docx."""

    def write_docx(self, path: Path, text: str) -> Path:
        try:
            from docx import Document  # type: ignore
        except ImportError as exc:
            raise ResumeParsingError("python-docx not installed") from exc
        path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        for line in text.splitlines():
            doc.add_paragraph(line)
        doc.save(str(path))
        return path
